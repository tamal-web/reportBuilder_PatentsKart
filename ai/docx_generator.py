# =============================================================================
# docx_generator.py — generate_docx(data, template, output_path)
# =============================================================================
"""
Converts a ReportData object (output of run_pipeline) into a formatted .docx
report with three sections:

  1. Summary Table      — one row per relevant patent
  2. Claim Charts       — one table per patent × all key features
  3. Novelty Matrix     — features (rows) × patents (columns), Yes/No cells

Call:
    generate_docx(report_data, template=TEMPLATES["default"], output_path="report.docx")
"""
from __future__ import annotations

import datetime
import logging
from pathlib import Path
from typing import List, Tuple

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import config as cfg
from config import ReportTemplate
from models import ReportData

log = logging.getLogger(__name__)


# =============================================================================
# Low-level XML helpers for python-docx
# =============================================================================


def _hex_to_rgb(hex_color: str) -> RGBColor:
    """'1B3A6B' → RGBColor(27, 58, 107)"""
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_bg(cell, hex_color: str) -> None:
    """Fill a table cell's background with a solid hex colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#").upper())
    # Remove any existing shd first
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    tcPr.append(shd)


def _set_table_borders(table, hex_color: str = "CBD5E1") -> None:
    """Apply single-line borders to every cell in a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr")) or OxmlElement("w:tblPr")
    tblBord = OxmlElement("w:tblBorders")
    color = hex_color.lstrip("#").upper()
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tblBord.append(el)
    # Remove any existing tblBorders
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    tblPr.append(tblBord)
    if tbl.find(qn("w:tblPr")) is None:
        tbl.insert(0, tblPr)


def _set_col_widths(table, widths_inches: List[float]) -> None:
    """Set individual column widths in inches."""
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_inches):
                cell.width = Inches(widths_inches[i])


def _cell_text(
    cell,
    text: str,
    t: ReportTemplate,
    bold: bool = False,
    center: bool = False,
    size_pt: int | None = None,
) -> None:
    """Write text into a cell, clearing any existing paragraphs first."""
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run = para.add_run(text)
    run.font.bold = bold
    run.font.name = t.font_name
    run.font.size = Pt(size_pt or t.font_size_pt)


def _header_cell(cell, text: str, t: ReportTemplate, center: bool = True) -> None:
    """Apply header styling (coloured background, white bold text) to a cell."""
    _set_cell_bg(cell, t.header_bg_color)
    _cell_text(cell, text, t, bold=True, center=center, size_pt=t.header_font_size_pt)
    run = cell.paragraphs[0].runs[0]
    run.font.color.rgb = _hex_to_rgb(t.header_text_color)


def _style_data_row(row, t: ReportTemplate, alt: bool = False) -> None:
    """Optionally apply alternating row background."""
    if alt:
        for cell in row.cells:
            _set_cell_bg(cell, t.alt_row_bg_color)


# =============================================================================
# Table builders
# =============================================================================

# Content width for US Letter with 1" margins: 8.5 - 2.0 = 6.5"
_PAGE_W = 6.5


def _build_summary_table(doc: Document, data: ReportData, t: ReportTemplate) -> None:
    """
    Summary Table
    ┌─────────────────────────┬──────────────────────┬──────────────┐
    │ Title                   │ Publication Number   │ Owner        │
    ├─────────────────────────┼──────────────────────┼──────────────┤
    │ ...                     │ ...                  │ ...          │
    └─────────────────────────┴──────────────────────┴──────────────┘
    """
    headers = ["Title", "Publication Number", "Owner"]
    col_widths = [3.0, 1.8, 1.7]  # sums to 6.5"

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    _set_table_borders(table, t.border_color)
    _set_col_widths(table, col_widths)

    for i, h in enumerate(headers):
        _header_cell(table.rows[0].cells[i], h, t)

    for idx, entry in enumerate(data.summary_table):
        row = table.add_row()
        _style_data_row(row, t, alt=(idx % 2 == 1))
        _cell_text(row.cells[0], entry.title, t)
        _cell_text(row.cells[1], entry.publication_number, t, center=True)
        _cell_text(row.cells[2], entry.owner, t)

    doc.add_paragraph()


def _build_claim_chart(doc: Document, data: ReportData, t: ReportTemplate) -> None:
    """
    Claim Chart — one table per patent
    ┌──────┬──────────────────────────────┬──────────────────────────────────┐
    │ No.  │ Key Feature                  │ Justification from Patent        │
    ├──────┼──────────────────────────────┼──────────────────────────────────┤
    │  1   │ Dual-layer encryption…       │ The patent discloses a method…   │
    └──────┴──────────────────────────────┴──────────────────────────────────┘
    """
    headers = ["No.", "Key Feature", "Justification from Patent"]
    col_widths = [0.45, 2.05, 4.0]  # sums to 6.5"

    for patent in data.patents:
        entries = sorted(
            data.claim_charts.get(patent.id, []),
            key=lambda e: e.feature_index,
        )

        # Subheading per patent
        heading = doc.add_heading(
            f"{patent.publication_number}  –  {patent.title}", level=2
        )
        heading.runs[0].font.size = Pt(12)

        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        _set_table_borders(table, t.border_color)
        _set_col_widths(table, col_widths)

        for i, h in enumerate(headers):
            _header_cell(table.rows[0].cells[i], h, t, center=(i == 0))

        for idx, entry in enumerate(entries):
            row = table.add_row()
            _style_data_row(row, t, alt=(idx % 2 == 1))
            _cell_text(row.cells[0], str(entry.feature_index), t, center=True)
            _cell_text(row.cells[1], entry.feature_description, t)
            _cell_text(row.cells[2], entry.justification, t)

        doc.add_paragraph()


def _build_matrix(doc: Document, data: ReportData, t: ReportTemplate) -> None:
    """
    Novelty Matrix
    ┌──────────────────────────┬──────────┬──────────┬──────────┐
    │ Key Feature              │ Patent 1 │ Patent 2 │ Patent 3 │
    ├──────────────────────────┼──────────┼──────────┼──────────┤
    │ Feature 1: Description…  │   Yes    │    No    │   Yes    │
    └──────────────────────────┴──────────┴──────────┴──────────┘
    """
    n_patents = len(data.matrix)
    if n_patents == 0:
        doc.add_paragraph("(No matrix data available.)")
        return

    n_cols = 1 + n_patents

    # Column widths: 2.5" for feature column, remainder split equally among patents
    feature_col_w = 2.5
    patent_col_w = round((_PAGE_W - feature_col_w) / n_patents, 4)
    col_widths = [feature_col_w] + [patent_col_w] * n_patents

    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    _set_table_borders(table, t.border_color)
    _set_col_widths(table, col_widths)

    # Header row: "Key Feature" + one column per patent
    hrow = table.rows[0]
    _header_cell(hrow.cells[0], "Key Feature", t, center=False)
    for i, m_entry in enumerate(data.matrix):
        label = (
            f"{m_entry.publication_number}\n{m_entry.patent_title[:30]}…"
            if len(m_entry.patent_title) > 30
            else f"{m_entry.publication_number}\n{m_entry.patent_title}"
        )
        _header_cell(hrow.cells[i + 1], label, t, center=True)

    # Data rows: one per key feature
    for idx, feature in enumerate(data.key_features):
        row = table.add_row()
        feature_label = f"Feature {feature.index}: {feature.description}"
        _cell_text(row.cells[0], feature_label, t)

        for i, m_entry in enumerate(data.matrix):
            found = m_entry.feature_results.get(feature.index, False)
            label = "Yes" if found else "No"
            cell = row.cells[i + 1]
            _set_cell_bg(cell, t.yes_cell_color if found else t.no_cell_color)
            _cell_text(cell, label, t, bold=True, center=True)

    doc.add_paragraph()


# =============================================================================
# Public API
# =============================================================================


def generate_docx(
    data: ReportData,
    template: ReportTemplate = None,
    output_path: str | Path = "prior_art_report.docx",
) -> Path:
    """
    Convert a fully populated ReportData object into a formatted .docx file.

    Args:
        data:        Output of run_pipeline().
        template:    ReportTemplate instance (defaults to cfg.TEMPLATES["default"]).
        output_path: Destination file path.

    Returns:
        The resolved output path (Path object).
    """
    if template is None:
        template = cfg.TEMPLATES["default"]

    output_path = Path(output_path)
    doc = Document()

    # ── Page setup (US Letter, 1" margins) ───────────────────────────────────
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_heading("Prior Art Search Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    try:
        logo_path = cfg.get_report_logo_path()
        if logo_path and logo_path.exists():
            doc.add_picture(str(logo_path), width=Inches(3))
        else:
            log.info("No logo image configured or file not found; skipping logo in docx.")
    except Exception as exc:
        log.warning(f"Could not insert logo image into docx: {exc}")

    meta = doc.add_paragraph(
        f"Generated: {datetime.datetime.now().strftime('%B %d, %Y')}    "
        f"Patents reviewed: {len(data.patents)}    "
        f"Key features: {len(data.key_features)}"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = _hex_to_rgb("64748B")

    doc.add_paragraph()

    # ─── Section 1: Summary Table ─────────────────────────────────────────────
    doc.add_heading("1.  Summary Table", level=1)
    _build_summary_table(doc, data, template)

    doc.add_page_break()

    # ─── Section 2: Claim Charts ──────────────────────────────────────────────
    doc.add_heading("2.  Claim Charts", level=1)
    _build_claim_chart(doc, data, template)

    doc.add_page_break()

    # ─── Section 3: Novelty Matrix ────────────────────────────────────────────
    doc.add_heading("3.  Novelty Matrix", level=1)
    p = doc.add_paragraph(
        "Yes = feature is disclosed in this patent  |  "
        "No = feature is not found in this patent"
    )
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.italic = True
    doc.add_paragraph()

    _build_matrix(doc, data, template)

    # ─── Save ─────────────────────────────────────────────────────────────────
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    log.info("Report saved → %s", output_path.resolve())
    print(f"✓ Report saved: {output_path.resolve()}")
    return output_path
